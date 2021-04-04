from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Inches

def true(var):
	if var:
		return 'Completed'
	else:
		return 'Not Completed'

def present(var):
	if var:
		return var
	else:
		return "Incomplete"

def create_document(standard,critique,samples,verifications):
	document = Document()

	document.add_heading('INTERNAL MODERATON COVER SHEET %s' %standard.year)
	paragraph = document.add_paragraph('NZQA Assessment (including Examination) Rules for Schools with Consent to Assess 2015 (5.6b): requires that schools must report to NZQA only those internal assessment results which have been subject to an internal moderation process.')

	document.add_heading('Standard Information', level=4)

	table1 = document.add_table(rows=9,cols=2,style='Table Grid')

	first_cells = table1.columns[0].cells

	first_cells[0].text = 'Course'
	first_cells[1].text = 'Standard Type'
	first_cells[2].text = 'Standard Number'
	first_cells[3].text = 'Version'
	first_cells[4].text = 'NZQF Level'
	first_cells[5].text = 'Credits'
	first_cells[6].text = 'Title'
	first_cells[7].text = 'Teacher in charge of assessment'

	second_cells = table1.columns[1].cells

	second_cells[0].text = standard.subject_title +' '+ standard.subject_id
	second_cells[1].text = present(standard.standard_type)
	second_cells[2].text = present(standard.standard_no)
	second_cells[3].text = present(standard.version)
	second_cells[4].text = present(standard.level)
	second_cells[5].text = present(str(standard.credits))
	second_cells[6].text = present(standard.title)
	second_cells[7].text = present(standard.tic)


	if critique:

		document.add_heading('CRITIQUING ASSESSMENT MATERIALS', level=4)
		paragraph = document.add_paragraph('The critiquing process ensures the assessment activity meets the requirements specified in the standard and provides the opportunity for students to present evidence at all grades. Assessment materials should be checked against the current clarification of the standard, conditions of assessment and/or any external moderation feedback before use.')

		table2 = document.add_table(rows=14,cols=2,style='Table Grid')

		tb2_first_cells = table2.columns[0].cells

		tb2_first_cells[0].text = 'Source of Materials'
		tb2_first_cells[1].text = 'Critiquer Name'
		tb2_first_cells[2].text = 'Critiquer School'
		tb2_first_cells[3].text = 'Date'
		tb2_first_cells[4].text = ''
		tb2_first_cells[5].text = 'The assessment material has been reviewed against the current standard clarification and/or external moderation feedback. (Where the material has been previously critiqued and the standard is unchanged, no further critiquing is required)'
		tb2_first_cells[6].text = 'Student instructions contain registered standard number, version, title, level & credits'
		tb2_first_cells[7].text = 'Student instructions are clear and language is appropriate'
		tb2_first_cells[8].text = 'The assessment is consistent with learning/context/curriculum at the appropriate level'
		tb2_first_cells[9].text = 'The assessment allows students to achieve all requirements of the standard for all grades'
		tb2_first_cells[10].text = 'Instructions are consistent with explanatory notes/range statements in the standard'
		tb2_first_cells[11].text = 'Assessment schedule is consistent with the standard and clarifications documents'
		tb2_first_cells[12].text = 'Judgement/ sufficiency statement clearly describe performance levels for each grade, e.g. quality & length'
		tb2_first_cells[13].text = 'Evidence statements allow for a range of acceptable answers from students with specific examples for each grade (N/A/M/E)'


		tb2_second_cells = table2.columns[1].cells

		tb2_second_cells[0].text = present(critique.materials)
		tb2_second_cells[1].text = present(critique.name)
		tb2_second_cells[2].text = present(critique.school)
		tb2_second_cells[3].text = present(critique.modified.strftime('%a %d %b at %I %p'))
		tb2_second_cells[4].text = ''
		tb2_second_cells[5].text = true(critique.check1)
		tb2_second_cells[6].text = true(critique.check2)
		tb2_second_cells[7].text = true(critique.check3)
		tb2_second_cells[8].text = true(critique.check4)
		tb2_second_cells[9].text = true(critique.check5)
		tb2_second_cells[10].text = true(critique.check6)
		tb2_second_cells[11].text = true(critique.check7)
		tb2_second_cells[12].text = true(critique.check8)
		tb2_second_cells[13].text = true(critique.check9)

		merged = tb2_first_cells[4].merge(tb2_second_cells[4])
		merged.add_paragraph('Critiqueing Process', style='Heading 4')

	if samples:

		document.add_heading('SAMPLES ARE RETAINED & ASSESSMENT MATERIALS REVIEWED', level=4)


		table3 = document.add_table(rows=12,cols=2,style='Table Grid')

		tb3_first_cells = table3.columns[0].cells

		tb3_first_cells[0].text = 'Teacher in Charge'
		tb3_first_cells[1].text = 'School of Teacher in Charge'
		tb3_first_cells[2].text = "The school's random selection procedure has been applied to select work for external moderation, if required."
		tb3_first_cells[3].text = 'Assessment materials have been reviewed in response to the assessor and/or verifier feedback.'
		tb3_first_cells[4].text = 'New benchmark samples have been annotated and/or existing examples of grade boundary decisions have been updated.'
		tb3_first_cells[5].text = 'Assessment materials and student work are available for external moderation'
		tb3_first_cells[6].text = 'Reviewed assessment materials are ready for future use.  Date:'
		tb3_first_cells[7].text = ''
		tb3_first_cells[8].text = 'Location Samples URL'
		tb3_first_cells[9].text = 'Physical Location of samples'
		tb3_first_cells[10].text = 'Location External Moderation URL'
		tb3_first_cells[11].text = 'Physical Location of moderation'


		tb3_second_cells = table3.columns[1].cells

		tb3_second_cells[0].text = present(samples.name)
		tb3_second_cells[1].text = present(samples.school)
		tb3_second_cells[2].text = true(samples.check1)
		tb3_second_cells[3].text = true(samples.check2)
		tb3_second_cells[4].text = true(samples.check3)
		tb3_second_cells[5].text = true(samples.check4)
		tb3_second_cells[6].text = true(samples.check5)
		tb3_second_cells[7].text = ''
		tb3_second_cells[8].text = present(samples.samples_url)
		tb3_second_cells[9].text = present(samples.samples_other)
		tb3_second_cells[10].text = present(samples.location_url)
		tb3_second_cells[11].text = present(samples.location_other)


		merged = tb3_first_cells[7].merge(tb2_second_cells[7])
		merged.add_paragraph('Locations', style='Heading 4')

	if verifications:

		document.add_heading('Verification Evidence', level=4)
		paragraph = document.add_paragraph('No of Verifications = %s' %len(verifications))

		for verification in verifications:
			table4 = document.add_table(rows=9,cols=2,style='Table Grid')		

			tb4_first_cells = table4.columns[0].cells

			tb4_first_cells[0].text = 'Student Name'
			tb4_first_cells[1].text = 'Markers Grade'
			tb4_first_cells[2].text = 'Verifiers Grade'
			tb4_first_cells[3].text = 'Reported Grade'
			tb4_first_cells[4].text = 'Discussion'
			tb4_first_cells[5].text = 'Verifiers Name'
			tb4_first_cells[6].text = 'Verifiers School'
			tb4_first_cells[7].text = 'Teacher in Charge'
			tb4_first_cells[8].text = ''

			tb4_second_cells = table4.columns[1].cells

			tb4_second_cells[0].text = present(verification.student)
			tb4_second_cells[1].text = present(verification.markers_grade)
			tb4_second_cells[2].text = present(verification.verifiers_grade)
			tb4_second_cells[3].text = present(verification.reported_grade)
			tb4_second_cells[4].text = present(verification.discussion)
			tb4_second_cells[5].text = present(verification.verifier_name)
			tb4_second_cells[6].text = present(verification.verifier_school)
			tb4_second_cells[7].text = present(verification.tic)
			tb4_second_cells[8].text = ''

			merged = tb4_first_cells[8].merge(tb4_second_cells[8])


	return document